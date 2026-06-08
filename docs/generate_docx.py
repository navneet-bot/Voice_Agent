"""Generate a single .docx file from all docs markdown files."""

import re
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from pathlib import Path

DOCS_DIR = Path(__file__).parent
OUTPUT_PATH = DOCS_DIR / "Cosmic_Chameleon_Project_Documentation.docx"


def parse_markdown(filepath):
    """Yield (type, content) tokens from a markdown file."""
    with open(filepath, encoding="utf-8") as f:
        lines = f.readlines()

    in_code_block = False
    code_buffer = []
    in_table = False
    table_buffer = []
    list_buffer = []  # (type, text) where type is 'ul' or 'ol'

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Code block
        if line.startswith("```"):
            if in_code_block:
                yield ("code", "\n".join(code_buffer))
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
                code_buffer = []
            i += 1
            continue
        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # Table (lines starting with |)
        if line.startswith("|") and line.endswith("|"):
            table_buffer.append(line)
            in_table = True
            i += 1
            continue
        else:
            if in_table and table_buffer:
                yield ("table", table_buffer)
                table_buffer = []
                in_table = False

        # Horizontal rule
        if re.match(r"^-{3,}$", line):
            yield ("hr", None)
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.+)$", line)
        if m:
            yield ("heading", (len(m.group(1)), m.group(2)))
            i += 1
            continue

        # Unordered list
        ulm = re.match(r"^(\s*)[-*+]\s+(.+)$", line)
        if ulm:
            yield ("ul", ulm.group(2))
            i += 1
            continue

        # Ordered list
        olm = re.match(r"^(\s*)\d+\.\s+(.+)$", line)
        if olm:
            yield ("ol", olm.group(2))
            i += 1
            continue

        # Bold line (like **text**)
        bm = re.match(r"^\*\*(.+)\*\*$", line)
        if bm:
            yield ("bold", bm.group(1))
            i += 1
            continue

        # Empty line
        if not line.strip():
            yield ("blank", None)
            i += 1
            continue

        # Regular paragraph
        yield ("para", line)
        i += 1

    # Flush buffers
    if in_code_block and code_buffer:
        yield ("code", "\n".join(code_buffer))
    if in_table and table_buffer:
        yield ("table", table_buffer)


def add_formatted_paragraph(doc, text, bold=False, italic=False, font_size=None, color=None):
    """Add a paragraph with inline formatting support."""
    p = doc.add_paragraph()
    # Parse inline formatting: **bold**, *italic*, `code`
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*|`.*?`)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            run = p.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
        else:
            run = p.add_run(part)
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        if font_size:
            run.font.size = Pt(font_size)
        if color:
            run.font.color.rgb = RGBColor(*color)
    return p


def add_table_from_lines(doc, table_lines):
    """Parse markdown table lines and add to document."""
    # Filter out separator rows
    data_rows = [l for l in table_lines if not re.match(r"^\|[-:\s|]+\|$", l)]
    if not data_rows:
        return

    rows = []
    for line in data_rows:
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)

    if not rows:
        return

    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    if i == 0:
                        run.bold = True


def add_code_block(doc, code_text):
    """Add a code block with monospace formatting."""
    for line in code_text.split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.3)
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def build_docx():
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)

    # Title page
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Cosmic Chameleon")
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Voice AI Calling SaaS Platform")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("Version 2.1 — Production Hardened\nJune 2026")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    doc.add_page_break()

    # Table of contents placeholder
    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = toc_title.add_run("Table of Contents")
    run.bold = True
    run.font.size = Pt(18)

    toc_items = [
        "1. Architecture",
        "2. Backend Modules",
        "3. Frontend Pages",
        "4. Voice Pipeline",
        "5. Database",
        "6. Deployment",
        "7. Testing",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(12)

    doc.add_page_break()

    # Files to process in order
    files = [
        ("Architecture", DOCS_DIR / "architecture.md"),
        ("Backend Modules", DOCS_DIR / "backend-modules.md"),
        ("Frontend Pages", DOCS_DIR / "frontend-pages.md"),
        ("Voice Pipeline", DOCS_DIR / "pipeline.md"),
        ("Database", DOCS_DIR / "database.md"),
        ("Deployment", DOCS_DIR / "deployment.md"),
        ("Testing", DOCS_DIR / "testing.md"),
    ]

    for section_title, md_file in files:
        if not md_file.exists():
            continue

        # Section title page
        p = doc.add_paragraph()
        run = p.add_run(section_title)
        run.bold = True
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
        doc.add_paragraph()

        tokens = list(parse_markdown(md_file))
        first_heading_seen = False

        for tok_type, tok_val in tokens:
            if tok_type == "heading":
                level, text = tok_val
                if not first_heading_seen and level == 1:
                    first_heading_seen = True
                    continue  # skip the doc title (already added as section title)
                heading = doc.add_heading(text, level=min(level, 4))
                for run in heading.runs:
                    run.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)

            elif tok_type == "para":
                add_formatted_paragraph(doc, tok_val)

            elif tok_type == "bold":
                p = doc.add_paragraph()
                run = p.add_run(tok_val)
                run.bold = True
                run.font.size = Pt(11)

            elif tok_type == "ul":
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(tok_val)
                run.font.size = Pt(10)

            elif tok_type == "ol":
                p = doc.add_paragraph(style="List Number")
                run = p.add_run(tok_val)
                run.font.size = Pt(10)

            elif tok_type == "table":
                add_table_from_lines(doc, tok_val)
                doc.add_paragraph()

            elif tok_type == "code":
                add_code_block(doc, tok_val)
                doc.add_paragraph()

            elif tok_type == "hr":
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                # Add a bottom border to simulate hr
                pPr = p._p.get_or_add_pPr()
                pBdr = pPr.makeelement(qn("w:pBdr"), {})
                bottom = pBdr.makeelement(qn("w:bottom"), {
                    qn("w:val"): "single",
                    qn("w:sz"): "6",
                    qn("w:space"): "1",
                    qn("w:color"): "999999",
                })
                pBdr.append(bottom)
                pPr.append(pBdr)

        doc.add_page_break()

    # Footer
    footer_text = "Cosmic Chameleon — Voice AI Calling SaaS Platform  |  Version 2.1  |  Confidential"
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(footer_text)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run.italic = True

    doc.save(OUTPUT_PATH)
    print(f"Document generated: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_docx()
